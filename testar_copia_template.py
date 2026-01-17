"""
Solução definitiva: Copiar o template como ZIP e trabalhar com ele.
Isso preserva TUDO: imagens, relationships, formatação.
"""
import shutil
import zipfile
import os
from docx import Document

def copiar_template_preservando_tudo(caminho_template, caminho_temp='temp_template.docx'):
    """Copia o template preservando absolutamente tudo via ZIP."""
    # Simplesmente copia o arquivo
    shutil.copy2(caminho_template, caminho_temp)
    return caminho_temp

# Teste
template_original = 'template_saida/template_saida.docx'
template_copia = copiar_template_preservando_tudo(template_original)

# Carrega a cópia
doc = Document(template_copia)

print(f"Template copiado: {template_copia}")
print(f"Seções: {len(doc.sections)}")
print(f"Parágrafos no cabeçalho: {len(doc.sections[0].header.paragraphs)}")

# Verifica se tem imagem
header = doc.sections[0].header
tem_imagem = False
for para in header.paragraphs:
    for run in para.runs:
        for element in run._element:
            if element.tag.endswith('drawing'):
                tem_imagem = True
                print("OK - IMAGEM ENCONTRADA no cabecalho!")
                break

if not tem_imagem:
    print("ERRO - Imagem NAO encontrada")

print("\nRelationships:")
for rel_id, rel in header.part.rels.items():
    print(f"  {rel_id}: {rel.reltype}")
