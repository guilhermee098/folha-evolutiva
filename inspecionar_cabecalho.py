"""Script para inspecionar o cabeçalho do template e extrair informações sobre imagens."""
from docx import Document
from docx.oxml.ns import qn

caminho = 'template_saida/template_saida.docx'
doc = Document(caminho)

print("="*60)
print("ANÁLISE DO CABEÇALHO DO TEMPLATE")
print("="*60)

for i, section in enumerate(doc.sections):
    print(f"\n>>> SEÇÃO {i+1}")
    header = section.header
    
    print(f"\nParágrafos no cabeçalho: {len(header.paragraphs)}")
    for j, para in enumerate(header.paragraphs):
        print(f"\n  Parágrafo {j+1}:")
        print(f"    Texto: {para.text[:100]}")
        print(f"    Alinhamento: {para.alignment}")
        print(f"    Runs: {len(para.runs)}")
        
        # Verifica se tem imagens (drawing)
        for run in para.runs:
            # Procura por elementos de imagem
            for element in run._element:
                if element.tag.endswith('drawing'):
                    print(f"      IMAGEM ENCONTRADA em Run!")
                    # Tenta extrair informações da imagem
                    try:
                        inline = element.find('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline')
                        if inline is not None:
                            extent = inline.find('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}extent')
                            if extent is not None:
                                print(f"        Largura: {extent.get('cx')}")
                                print(f"        Altura: {extent.get('cy')}")
                    except:
                        pass
    
    # Verifica relationships (imagens)
    print(f"\n  Relationships no cabeçalho:")
    try:
        for rel_id, rel in header.part.rels.items():
            print(f"    {rel_id}: {rel.reltype}")
            if 'image' in rel.reltype.lower():
                print(f"      → IMAGEM: {rel.target_ref}")
    except:
        pass

print("\n" + "="*60)
