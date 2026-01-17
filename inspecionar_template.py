"""Script temporário para inspecionar a estrutura do documento modelo."""
from docx import Document
from docx.shared import Pt, Inches

caminho = 'saida/JOAO PAULO NUNES - FOLHA DE EVOLUÇÃO TRANSDISCIPLINA - DIARIA.docx'
doc = Document(caminho)

print("="*60)
print("ANALISE DO DOCUMENTO MODELO")
print("="*60)

# Informações gerais
print(f"\nParagrafos: {len(doc.paragraphs)}")
print(f"Tabelas: {len(doc.tables)}")
print(f"Secoes: {len(doc.sections)}")

# Margens
section = doc.sections[0]
print(f"\nMARGENS:")
print(f"  Superior: {section.top_margin.inches:.2f} polegadas")
print(f"  Inferior: {section.bottom_margin.inches:.2f} polegadas")
print(f"  Esquerda: {section.left_margin.inches:.2f} polegadas")
print(f"  Direita: {section.right_margin.inches:.2f} polegadas")

# Cabeçalho
print(f"\nCABECALHO:")
if section.header.paragraphs:
    for i, para in enumerate(section.header.paragraphs[:3]):
        texto = para.text[:80] if para.text else "(vazio)"
        print(f"  Paragrafo {i+1}: {texto}")

# Título principal
print(f"\nTITULO PRINCIPAL:")
if doc.paragraphs:
    primeiro = doc.paragraphs[0]
    print(f"  Texto: {primeiro.text}")
    if primeiro.runs:
        run = primeiro.runs[0]
        print(f"  Fonte: {run.font.name}")
        print(f"  Tamanho: {run.font.size.pt if run.font.size else 'padrao'} pt")
        print(f"  Negrito: {run.font.bold}")
        print(f"  Alinhamento: {primeiro.alignment}")

# Estrutura das tabelas
print(f"\nESTRUTURA DAS TABELAS:")
for i, tabela in enumerate(doc.tables[:3], 1):  # Primeiras 3 tabelas
    print(f"\n  Tabela {i}:")
    print(f"    Linhas: {len(tabela.rows)}")
    print(f"    Colunas: {len(tabela.columns)}")
    print(f"    Estilo: {tabela.style.name if tabela.style else 'Nenhum'}")
    
    # Cabeçalho da tabela
    if tabela.rows:
        cabecalho = [cell.text for cell in tabela.rows[0].cells]
        print(f"    Cabecalho: {cabecalho}")
        
        # Formatação da primeira célula do cabeçalho
        primeira_celula = tabela.rows[0].cells[0]
        if primeira_celula.paragraphs and primeira_celula.paragraphs[0].runs:
            run = primeira_celula.paragraphs[0].runs[0]
            print(f"    Fonte cabecalho: {run.font.name}")
            print(f"    Tamanho cabecalho: {run.font.size.pt if run.font.size else 'padrao'} pt")

# Títulos de seção (headings)
print(f"\nTITULOS DE SECAO:")
headings = [p for p in doc.paragraphs if p.style.name.startswith('Heading')]
for h in headings[:5]:
    print(f"  {h.style.name}: {h.text}")
    if h.runs:
        run = h.runs[0]
        print(f"    Fonte: {run.font.name}, Tamanho: {run.font.size.pt if run.font.size else 'padrao'} pt")

print("\n" + "="*60)
