"""
Script para criar/recriar o template base a partir de um documento existente.

Uso:
    python criar_template.py [caminho_do_documento.docx]

Se nenhum caminho for fornecido, usa o arquivo da pasta saida/ como padrão.
"""

from docx import Document
import sys
import os

def criar_template_de_documento(caminho_origem, caminho_destino='template_evolucao.docx'):
    """Cria um template a partir de um documento existente."""
    
    if not os.path.exists(caminho_origem):
        print(f"ERRO: Arquivo '{caminho_origem}' nao encontrado!")
        return False
    
    print(f"Carregando documento: '{caminho_origem}'")
    doc_original = Document(caminho_origem)
    
    # Cria novo documento
    doc_template = Document()
    
    print("Copiando configuracoes de pagina e margens...")
    # Copia as configurações de seção
    for section_orig in doc_original.sections:
        section_nova = doc_template.sections[0] if len(doc_template.sections) > 0 else doc_template.add_section()
        section_nova.top_margin = section_orig.top_margin
        section_nova.bottom_margin = section_orig.bottom_margin
        section_nova.left_margin = section_orig.left_margin
        section_nova.right_margin = section_orig.right_margin
        section_nova.page_height = section_orig.page_height
        section_nova.page_width = section_orig.page_width
        
        print("Copiando cabecalho...")
        # Copia o cabeçalho
        header_orig = section_orig.header
        header_nova = section_nova.header
        
        # Remove parágrafos padrão
        for para in header_nova.paragraphs:
            para.clear()
        
        for para in header_orig.paragraphs:
            new_para = header_nova.add_paragraph()
            for run in para.runs:
                new_run = new_para.add_run(run.text)
                new_run.bold = run.bold
                new_run.italic = run.italic
                new_run.underline = run.underline
                if run.font.size:
                    new_run.font.size = run.font.size
                if run.font.name:
                    new_run.font.name = run.font.name
            new_para.alignment = para.alignment
    
    # Adiciona título de exemplo
    print("Criando secao modelo (FISIOTERAPIA)...")
    titulo = doc_template.add_paragraph('FISIOTERAPIA')
    titulo.style = 'Heading 1'
    
    # Copia APENAS a primeira tabela como modelo
    if doc_original.tables:
        print("Copiando tabela modelo...")
        tabela_original = doc_original.tables[0]
        
        # Adiciona tabela ao template
        tabela_template = doc_template.add_table(rows=len(tabela_original.rows), cols=len(tabela_original.columns))
        tabela_template.style = tabela_original.style
        
        # Copia formatação de cada célula
        for i, row_orig in enumerate(tabela_original.rows):
            for j, cell_orig in enumerate(row_orig.cells):
                cell_template = tabela_template.rows[i].cells[j]
                
                # Limpa conteúdo
                for para in cell_template.paragraphs:
                    para.clear()
                
                # Copia formatação do cabeçalho ou deixa vazio
                if i == 0:  # Cabeçalho
                    for para_orig in cell_orig.paragraphs:
                        para_template = cell_template.paragraphs[0] if cell_template.paragraphs else cell_template.add_paragraph()
                        for run_orig in para_orig.runs:
                            run_template = para_template.add_run(run_orig.text)
                            run_template.bold = run_orig.bold
                            run_template.italic = run_orig.italic
                            if run_orig.font.size:
                                run_template.font.size = run_orig.font.size
                            if run_orig.font.name:
                                run_template.font.name = run_orig.font.name
                        para_template.alignment = para_orig.alignment
                else:
                    # Linhas de dados: deixa vazio mas mantém estrutura
                    cell_template.text = ""
    else:
        print("AVISO: Nenhuma tabela encontrada no documento!")
    
    # Salva o template
    print(f"Salvando template: '{caminho_destino}'")
    doc_template.save(caminho_destino)
    
    print("\nOK - Template criado com sucesso!")
    print(f"  - Arquivo: '{caminho_destino}'")
    print(f"  - Formatacao preservada")
    print(f"  - Cabecalho copiado")
    print(f"  - Uma tabela modelo")
    print(f"  - Margens: {section_nova.top_margin.inches:.2f}\" (topo)")
    
    return True

if __name__ == '__main__':
    # Determina o arquivo de origem
    if len(sys.argv) > 1:
        arquivo_origem = sys.argv[1]
    else:
        # Usa o arquivo padrão da pasta saida
        arquivo_origem = 'saida/JOAO PAULO NUNES - FOLHA DE EVOLUÇÃO TRANSDISCIPLINA - DIARIA.docx'
    
    print("="*60)
    print("CRIADOR DE TEMPLATE")
    print("="*60 + "\n")
    
    criar_template_de_documento(arquivo_origem)
