"""
Teste de 3 alternativas para adicionar logo ao cabeçalho.
"""
from docx import Document
from docx.shared import Inches
import shutil

# Copia o template
shutil.copy2('template_saida/template_saida.docx', 'teste_logo.docx')
doc = Document('teste_logo.docx')

header = doc.sections[0].header
para_titulo = header.paragraphs[0]

print("="*60)
print("TESTANDO 3 ALTERNATIVAS PARA ADICIONAR LOGO")
print("="*60)

# Verifica se já tem logo
tem_logo = False
for run in para_titulo.runs:
    for element in run._element:
        if element.tag.endswith('drawing'):
            tem_logo = True
            print("\nOK - LOGO JA EXISTE NO TEMPLATE")
            print(f"  Texto do paragrafo: {para_titulo.text[:80]}...")
            break

if tem_logo:
    print("\n>>> Logo encontrada! O template esta correto.")
    print(">>> O problema pode estar na substituicao de variaveis.")
else:
    print("\n>>> Logo NAO encontrada. Testando alternativas...")
    
    print("\n--- ALTERNATIVA 1: add_picture no primeiro run ---")
    try:
        if para_titulo.runs:
            para_titulo.runs[0].add_picture('logo_extraida.png', width=Inches(0.8))
            doc.save('teste_logo_alt1.docx')
            print("OK - Salvo como: teste_logo_alt1.docx")
    except Exception as e:
        print(f"ERRO: {e}")
    
    # Recarrega
    doc = Document('teste_logo.docx')
    header = doc.sections[0].header
    para_titulo = header.paragraphs[0]
    
    print("\n--- ALTERNATIVA 2: Novo run no inicio com imagem ---")
    try:
        # Insere run no início
        run_logo = para_titulo.insert_paragraph_before().add_run()
        run_logo.add_picture('logo_extraida.png', width=Inches(0.8))
        # Move de volta
        doc.save('teste_logo_alt2.docx')
        print("OK - Salvo como: teste_logo_alt2.docx")
    except Exception as e:
        print(f"ERRO: {e}")
    
    # Recarrega
    doc = Document('teste_logo.docx')
    header = doc.sections[0].header
    para_titulo = header.paragraphs[0]
    
    print("\n--- ALTERNATIVA 3: Criar paragrafo separado para logo ---")
    try:
        # Cria parágrafo antes do título
        para_logo = header.paragraphs[0].insert_paragraph_before()
        run_logo = para_logo.add_run()
        run_logo.add_picture('logo_extraida.png', width=Inches(0.8))
        doc.save('teste_logo_alt3.docx')
        print("OK - Salvo como: teste_logo_alt3.docx")
    except Exception as e:
        print(f"ERRO: {e}")

print("\n" + "="*60)
print("Abra os arquivos gerados para ver qual alternativa funcionou!")
print("="*60)
