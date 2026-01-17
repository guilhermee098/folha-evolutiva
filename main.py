from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from datetime import datetime, timedelta
import os
import json
import logging
import sys
import copy

# Configuração de logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    datefmt='%H:%M:%S'
)
logger = logging.getLogger(__name__)

# Configurações padrão
CONFIG_PADRAO = {
    "duracao_atendimento_minutos": 40,
    "colunas_esperadas": ["DATA", "HORÁRIO", "PROCEDIMENTO"],
    "formato_hora": "%H:%M",
    "permitir_data_vazia_primeira_linha": False,
    "caminho_template": "template_saida/template_saida.docx",
    "extrair_cabecalho_de_entrada": True
}

def carregar_configuracao(caminho_config='config.json'):
    """Carrega configurações do arquivo JSON ou usa padrões."""
    if os.path.exists(caminho_config):
        try:
            with open(caminho_config, 'r', encoding='utf-8') as f:
                config = json.load(f)
                logger.info(f"✓ Configurações carregadas de '{caminho_config}'")
                return {**CONFIG_PADRAO, **config}
        except json.JSONDecodeError as e:
            logger.warning(f"⚠ Erro ao ler config.json: {e}. Usando configurações padrão.")
            return CONFIG_PADRAO
        except Exception as e:
            logger.warning(f"⚠ Erro inesperado ao carregar config: {e}. Usando configurações padrão.")
            return CONFIG_PADRAO
    else:
        logger.info("ℹ Arquivo config.json não encontrado. Usando configurações padrão.")
        return CONFIG_PADRAO

def validar_arquivo_entrada(caminho):
    """Valida se o arquivo de entrada existe e é válido."""
    logger.info(f"→ Validando arquivo de entrada: '{caminho}'")
    
    if not os.path.exists(caminho):
        logger.error(f"✗ ERRO: Arquivo '{caminho}' não encontrado!")
        return False
    
    if not caminho.endswith('.docx'):
        logger.error(f"✗ ERRO: Arquivo '{caminho}' não é um documento Word (.docx)!")
        return False
    
    try:
        # Tenta abrir o documento para verificar se é válido
        Document(caminho)
        logger.info(f"✓ Arquivo válido e acessível")
        return True
    except Exception as e:
        logger.error(f"✗ ERRO: Não foi possível abrir o documento: {e}")
        return False

def extrair_dados_cabecalho(caminho_origem):
    """Extrai dados do cabeçalho do documento de entrada."""
    logger.info(f"→ Extraindo dados do cabeçalho")
    
    try:
        doc = Document(caminho_origem)
        
        # Tenta extrair do cabeçalho do documento
        if doc.sections and doc.sections[0].header.paragraphs:
            header_text = "\n".join([p.text for p in doc.sections[0].header.paragraphs])
            
            # Extração de dados (adaptável ao formato)
            dados_cabecalho = {
                "nome_paciente": "",
                "iniciais": "",
                "data_nascimento": "",
                "mes_ano": "",
                "codigos_cid_e_descricao": ""
            }
            
            # Busca por padrões no cabeçalho
            import re
            
            # Nome (sem iniciais neste formato): "Nome: João Paulo Braz Nunes"
            match_nome = re.search(r'Nome:\s*([^\n]+?)(?:\s+Nasc)', header_text, re.IGNORECASE)
            if match_nome:
                nome_completo = match_nome.group(1).strip()
                dados_cabecalho["nome_paciente"] = nome_completo
                
                # Gera iniciais automaticamente
                palavras = nome_completo.split()
                iniciais = ''.join([p[0].upper() for p in palavras if p])
                dados_cabecalho["iniciais"] = iniciais
            
            # Data de nascimento: "Nasc.: 27/12/2018" ou "Nasc: 27/12/2018"
            match_nasc = re.search(r'Nasc\.?:\s*(\d{2}/\d{2}/\d{4})', header_text, re.IGNORECASE)
            if match_nasc:
                dados_cabecalho["data_nascimento"] = match_nasc.group(1)
            
            # Mês/Ano: "MÊS DE JULHO/2025" ou "JULHO/2025"
            match_mes = re.search(r'(?:M[ÊE]S\s+DE\s+)?([A-Z]+/\d{4})', header_text, re.IGNORECASE)
            if match_mes:
                dados_cabecalho["mes_ano"] = match_mes.group(1).upper()
            
            # CID: Captura todos os CIDs e descrições em uma linha
            # Formato: "F84.9 Transtornos..., F90.0 Transtorno..., 6A02.Z Transtorno..."
            match_cid = re.search(r'Diagn[óo]stico:\s*(.+?)(?:\n|$)', header_text, re.IGNORECASE | re.DOTALL)
            if match_cid:
                diagnostico_completo = match_cid.group(1).strip()
                # Remove espaços extras e normaliza
                diagnostico_completo = ' '.join(diagnostico_completo.split())
                dados_cabecalho["codigos_cid_e_descricao"] = diagnostico_completo
            
            # Verifica se extraiu algum dado
            tem_dados = any(v for v in dados_cabecalho.values() if v)
            
            if tem_dados:
                logger.info(f"✓ Dados extraídos do cabeçalho:")
                if dados_cabecalho['nome_paciente']:
                    logger.info(f"  • Nome: {dados_cabecalho['nome_paciente']}")
                if dados_cabecalho['iniciais']:
                    logger.info(f"  • Iniciais: {dados_cabecalho['iniciais']}")
                if dados_cabecalho['data_nascimento']:
                    logger.info(f"  • Data Nascimento: {dados_cabecalho['data_nascimento']}")
                if dados_cabecalho['mes_ano']:
                    logger.info(f"  • Mês/Ano: {dados_cabecalho['mes_ano']}")
                if dados_cabecalho['codigos_cid_e_descricao']:
                    logger.info(f"  • CID: {dados_cabecalho['codigos_cid_e_descricao']}")
                
                return dados_cabecalho
            else:
                logger.warning("⚠ Nenhum dado extraído do cabeçalho (padrões não encontrados)")
                logger.info(f"ℹ Conteúdo do cabeçalho: {header_text[:200]}...")
                return None
            
        else:
            logger.warning("⚠ Cabeçalho não encontrado no documento de entrada")
            return None
            
    except Exception as e:
        logger.error(f"✗ Erro ao extrair dados do cabeçalho: {e}")
        return None

def identificar_e_extrair_tabelas(caminho_origem, config):
    """Identifica e extrai dados das tabelas do documento."""
    logger.info(f"→ Iniciando extração de dados de '{caminho_origem}'")
    
    doc = Document(caminho_origem)
    dados_totais = []
    tabelas_encontradas = 0
    erros_parsing = []
    avisos_data = []
    
    colunas_alvo = config["colunas_esperadas"]
    logger.info(f"→ Procurando tabelas com colunas: {', '.join(colunas_alvo)}")

    for idx_tabela, tabela in enumerate(doc.tables, start=1):
        # Verifica se a primeira linha da tabela contém nossas colunas alvo
        cabecalho = [celula.text.strip().upper() for celula in tabela.rows[0].cells]
        
        if all(coluna in cabecalho for coluna in colunas_alvo):
            tabelas_encontradas += 1
            logger.info(f"✓ Tabela {idx_tabela} identificada como válida ({len(tabela.rows)-1} linhas de dados)")
            
            # Mapeia os índices das colunas
            idx_data = cabecalho.index("DATA")
            idx_hora = cabecalho.index("HORÁRIO")
            idx_proc = cabecalho.index("PROCEDIMENTO")
            
            data_atual = ""
            for i, row in enumerate(tabela.rows):
                if i == 0: continue # Pula o cabeçalho
                
                texto_data = row.cells[idx_data].text.strip()
                texto_hora = row.cells[idx_hora].text.strip()
                texto_proc = row.cells[idx_proc].text.strip().upper()

                # Lógica de persistência da data
                if texto_data:
                    data_atual = texto_data
                
                # Validação: primeira linha sem data
                if not data_atual and texto_hora and texto_proc:
                    if not config["permitir_data_vazia_primeira_linha"]:
                        msg = f"Tabela {idx_tabela}, Linha {i+1}: Registro sem data ('{texto_proc}' às {texto_hora})"
                        avisos_data.append(msg)
                        logger.warning(f"⚠ {msg}")
                        continue
                
                if texto_hora and texto_proc:
                    # Validação de formato de hora
                    try:
                        datetime.strptime(texto_hora, config["formato_hora"])
                        dados_totais.append({
                            "data": data_atual,
                            "inicio": texto_hora,
                            "procedimento": texto_proc,
                            "linha_origem": i + 1,
                            "tabela_origem": idx_tabela
                        })
                    except ValueError:
                        msg = f"Tabela {idx_tabela}, Linha {i+1}: Formato de hora inválido '{texto_hora}' (esperado {config['formato_hora']})"
                        erros_parsing.append(msg)
                        logger.warning(f"⚠ {msg}")
    
    # Relatório de extração
    logger.info(f"\n{'='*60}")
    logger.info(f"RESUMO DA EXTRAÇÃO:")
    logger.info(f"  • Tabelas válidas encontradas: {tabelas_encontradas}")
    logger.info(f"  • Registros extraídos: {len(dados_totais)}")
    
    if erros_parsing:
        logger.warning(f"  • Erros de parsing: {len(erros_parsing)}")
    if avisos_data:
        logger.warning(f"  • Registros sem data: {len(avisos_data)}")
    
    logger.info(f"{'='*60}\n")
    
    if tabelas_encontradas == 0:
        logger.error(f"✗ ERRO: Nenhuma tabela com as colunas esperadas foi encontrada!")
        logger.error(f"   Certifique-se de que o documento contém uma tabela com: {', '.join(colunas_alvo)}")
    
    return dados_totais, erros_parsing, avisos_data

def adicionar_logo_e_confidencial_ao_cabecalho(doc, caminho_logo='logo_extraida.png'):
    """Adiciona a logo e o texto CONFIDENCIAL ao cabeçalho do documento."""
    from docx.shared import Inches, RGBColor, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    logger.info("→ Verificando logo e CONFIDENCIAL no cabeçalho")
    
    try:
        for section in doc.sections:
            header = section.header
            
            # Procura pelo primeiro parágrafo (que contém o título)
            if not header.paragraphs:
                logger.warning("⚠ Cabeçalho sem parágrafos")
                continue
            
            para_titulo = header.paragraphs[0]
            
            # Verifica se já tem a logo
            tem_imagem = False
            for run in para_titulo.runs:
                for element in run._element:
                    if element.tag.endswith('drawing'):
                        tem_imagem = True
                        break
            
            if tem_imagem:
                logger.info("✓ Logo já está presente no cabeçalho")
            else:
                logger.warning("⚠ Logo não encontrada no cabeçalho")
                logger.info("→ Tentando adicionar logo manualmente...")
                
                # Tenta adicionar a logo do arquivo extraído
                if os.path.exists(caminho_logo):
                    # Cria um novo run NO INÍCIO do parágrafo para a imagem
                    # Precisamos inserir ANTES do primeiro run existente
                    if para_titulo.runs:
                        # Pega a formatação do primeiro run
                        primeiro_run = para_titulo.runs[0]
                        
                        # Cria um novo run com a imagem
                        run_logo = para_titulo.insert_paragraph_before().add_run()
                        run_logo.add_picture(caminho_logo, width=Inches(0.8))
                        
                        # Move o parágrafo com imagem para dentro do header na posição correta
                        # Na verdade, vamos adicionar a imagem diretamente no primeiro run
                        try:
                            # Tenta adicionar a imagem no início do primeiro run
                            from docx.oxml import parse_xml
                            from docx.oxml.ns import nsdecls
                            
                            # Adiciona a imagem
                            primeiro_run.add_picture(caminho_logo, width=Inches(0.8))
                            logger.info(f"✓ Logo adicionada manualmente ao cabeçalho")
                        except Exception as e:
                            logger.warning(f"⚠ Não foi possível adicionar logo: {e}")
            
            # Adiciona CONFIDENCIAL no mesmo parágrafo, mas à direita
            texto_completo = para_titulo.text
            
            if 'CONFIDENCIAL' not in texto_completo:
                # Adiciona tab stop à direita
                from docx.enum.text import WD_TAB_ALIGNMENT
                tab_stops = para_titulo.paragraph_format.tab_stops
                tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
                
                # Encontra o último run do título para copiar sua formatação
                ultimo_run_titulo = para_titulo.runs[-1] if para_titulo.runs else None
                
                # Adiciona um run com tab (SEM formatação especial - usa a do parágrafo)
                run_tab = para_titulo.add_run('\t\t')
                # Copia a formatação do último run do título
                if ultimo_run_titulo:
                    run_tab.font.name = ultimo_run_titulo.font.name
                    run_tab.font.size = ultimo_run_titulo.font.size
                    run_tab.font.bold = ultimo_run_titulo.font.bold
                
                # Adiciona run com CONFIDENCIAL (com formatação própria)
                run_conf = para_titulo.add_run('CONFIDENCIAL')
                run_conf.font.color.rgb = RGBColor(255, 0, 0)  # Vermelho
                run_conf.font.bold = True
                run_conf.font.size = Pt(14)
                
                logger.info("✓ Texto CONFIDENCIAL adicionado à direita do título")
            else:
                logger.info("✓ Texto CONFIDENCIAL já está presente")
        
        return True
    except Exception as e:
        logger.error(f"✗ Erro ao verificar logo/CONFIDENCIAL: {e}")
        import traceback
        traceback.print_exc()
        return False

def copiar_cabecalho_completo(doc_template, doc_destino):
    """
    Copia o cabeçalho completo do template para o documento de destino,
    incluindo imagens, formatação, cores, etc.
    """
    from copy import deepcopy
    
    logger.info("→ Copiando cabeçalho completo do template")
    
    # Copia cada seção
    for i, section_template in enumerate(doc_template.sections):
        if i < len(doc_destino.sections):
            section_destino = doc_destino.sections[i]
            
            # Remove o cabeçalho atual do destino
            header_destino = section_destino.header
            for elemento in list(header_destino._element):
                header_destino._element.remove(elemento)
            
            # Clona todos os elementos do cabeçalho do template
            header_template = section_template.header
            for elemento in header_template._element:
                elemento_clonado = deepcopy(elemento)
                header_destino._element.append(elemento_clonado)
            
            logger.info(f"✓ Cabeçalho da seção {i+1} copiado com sucesso")

def substituir_variaveis_cabecalho(doc, dados_cabecalho):
    """Substitui as variáveis no cabeçalho do documento preservando TODA a formatação."""
    if not dados_cabecalho:
        logger.warning("⚠ Nenhum dado de cabeçalho para substituir")
        return
    
    logger.info("→ Substituindo variáveis no cabeçalho do template")
    
    # Mapeamento de variáveis
    variaveis = {
        "{MES_ANO}": dados_cabecalho.get("mes_ano", ""),
        "{NOME_PACIENTE}": dados_cabecalho.get("nome_paciente", ""),
        "{INICIAIS}": dados_cabecalho.get("iniciais", ""),
        "{DATA_NASCIMENTO}": dados_cabecalho.get("data_nascimento", ""),
        "{CODIGOS_CID_E_DESCRICAO}": dados_cabecalho.get("codigos_cid_e_descricao", "")
    }
    
    # Substitui no cabeçalho
    for section in doc.sections:
        for para in section.header.paragraphs:
            substituir_variaveis_em_paragrafo(para, variaveis)
    
    # Também substitui no corpo (caso tenha variáveis lá)
    for para in doc.paragraphs:
        substituir_variaveis_em_paragrafo(para, variaveis)

def substituir_variaveis_em_paragrafo(para, variaveis):
    """Substitui variáveis em um parágrafo preservando formatação de cada run E IMAGENS."""
    # Reconstrói o texto completo
    texto_completo = ''.join([run.text for run in para.runs])
    
    # Verifica se há variáveis
    tem_variaveis = any(var in texto_completo for var in variaveis.keys())
    
    if not tem_variaveis:
        return
    
    # Substitui as variáveis no texto
    texto_novo = texto_completo
    for var, valor in variaveis.items():
        if var in texto_novo:
            texto_novo = texto_novo.replace(var, valor)
            logger.info(f"  ✓ {var} → {valor[:50]}..." if len(valor) > 50 else f"  ✓ {var} → {valor}")
    
    # IMPORTANTE: Preservar runs com imagens (drawing elements)
    runs_com_imagem = []
    for i, run in enumerate(para.runs):
        for element in run._element:
            if element.tag.endswith('drawing'):
                runs_com_imagem.append(i)
                break
    
    # Se tem runs com imagem, NÃO redistribuir o texto (pode perder a imagem)
    # Em vez disso, apenas substituir variáveis run por run
    if runs_com_imagem:
        logger.info("  ⚠ Parágrafo contém imagens - substituindo variáveis com cuidado")
        for run in para.runs:
            # Pula runs com imagens
            tem_imagem_neste_run = False
            for element in run._element:
                if element.tag.endswith('drawing'):
                    tem_imagem_neste_run = True
                    break
            
            if not tem_imagem_neste_run:
                texto_run = run.text
                for var, valor in variaveis.items():
                    if var in texto_run:
                        texto_run = texto_run.replace(var, valor)
                run.text = texto_run
        return
    
    # Se não tem imagens, pode redistribuir normalmente
    tamanho_original = len(texto_completo)
    tamanho_novo = len(texto_novo)
    
    if tamanho_original == 0:
        return
    
    # Distribui o novo texto pelos runs existentes mantendo proporções
    posicao_atual = 0
    for i, run in enumerate(para.runs):
        tamanho_run_original = len(run.text)
        
        if tamanho_run_original == 0:
            run.text = ""
            continue
        
        # Calcula a proporção deste run no texto total
        proporcao = tamanho_run_original / tamanho_original
        tamanho_run_novo = int(proporcao * tamanho_novo)
        
        # Último run pega o resto
        if i == len(para.runs) - 1:
            run.text = texto_novo[posicao_atual:]
        else:
            run.text = texto_novo[posicao_atual:posicao_atual + tamanho_run_novo]
            posicao_atual += tamanho_run_novo

def clonar_tabela_completa(tabela_modelo, doc):
    """
    Clona uma tabela preservando TODAS as propriedades:
    - Larguras de colunas
    - Alturas de linhas
    - Estilo da tabela
    - Bordas, cores, fonte, alinhamento
    """
    from docx.oxml import parse_xml
    from copy import deepcopy
    
    # Clona o elemento XML da tabela inteira
    tbl_element = deepcopy(tabela_modelo._element)
    
    # Adiciona a tabela clonada ao documento
    doc._element.body.append(tbl_element)
    
    # Retorna a nova tabela (último item na lista de tabelas)
    return doc.tables[-1]

def clonar_linha_tabela(tabela_destino, linha_modelo):
    """Clona uma linha de tabela preservando toda a formatação."""
    from copy import deepcopy
    
    # Clona o elemento XML da linha
    tr_element = deepcopy(linha_modelo._element)
    
    # Adiciona a linha clonada à tabela
    tabela_destino._element.append(tr_element)
    
    # Retorna a nova linha
    return tabela_destino.rows[-1]

def preencher_linha_tabela(linha, dados_linha):
    """Preenche uma linha da tabela com dados mantendo a formatação."""
    for i, dado in enumerate(dados_linha):
        if i < len(linha.cells):
            celula = linha.cells[i]
            # Limpa o conteúdo preservando formatação
            if celula.paragraphs:
                para = celula.paragraphs[0]
                # Limpa texto mas mantém formatação dos runs
                for run in para.runs:
                    run.text = ''
                # Adiciona o novo texto no primeiro run (mantém formatação)
                if para.runs:
                    para.runs[0].text = str(dado)
                else:
                    # Se não tem runs, adiciona texto simples
                    para.add_run(str(dado))

def gerar_word_evolucao(dados, caminho_destino, config, dados_cabecalho=None):
    """Gera o documento Word de evolução usando template com substituição de variáveis."""
    if not dados:
        logger.error("✗ Nenhum dado válido para gerar o documento de evolução.")
        return False

    logger.info(f"→ Gerando documento de evolução: '{caminho_destino}'")
    
    # Verifica se o template existe
    caminho_template = config["caminho_template"]
    if not os.path.exists(caminho_template):
        logger.error(f"✗ Template não encontrado: '{caminho_template}'")
        logger.error("   Certifique-se de que o template foi criado corretamente.")
        return False
    
    logger.info(f"→ Carregando template: '{caminho_template}'")
    
    # SOLUÇÃO DEFINITIVA: Copiar o template fisicamente antes de abrir
    # Isso preserva TUDO: imagens, relationships, formatação
    import shutil
    import tempfile
    
    # Cria uma cópia temporária do template
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    temp_template_path = temp_file.name
    temp_file.close()
    
    shutil.copy2(caminho_template, temp_template_path)
    logger.info(f"✓ Template copiado para: {temp_template_path}")
    
    # Carrega a CÓPIA do template (que preserva tudo)
    doc = Document(temp_template_path)
    
    logger.info("✓ Template carregado (cabeçalho preservado com logo e CONFIDENCIAL)")
    
    # Adiciona logo e CONFIDENCIAL se não estiverem presentes
    adicionar_logo_e_confidencial_ao_cabecalho(doc)
    
    # Substitui variáveis do cabeçalho (SEM remover/recriar nada)
    if dados_cabecalho:
        substituir_variaveis_cabecalho(doc, dados_cabecalho)
    
    # Localiza a tabela modelo (primeira tabela do template)
    if not doc.tables:
        logger.error("✗ Template não contém nenhuma tabela modelo!")
        return False
    
    tabela_modelo = doc.tables[0]
    logger.info(f"✓ Tabela modelo identificada: {len(tabela_modelo.rows)} linhas, {len(tabela_modelo.columns)} colunas")
    
    # Verifica se a tabela modelo tem pelo menos 2 linhas (cabeçalho + 1 linha de exemplo)
    if len(tabela_modelo.rows) < 2:
        logger.error("✗ Tabela modelo deve ter pelo menos 2 linhas (cabeçalho + linha de dados exemplo)")
        return False
    
    linha_dados_modelo = tabela_modelo.rows[1]  # Segunda linha como modelo de formatação
    
    # Localiza e salva o parágrafo do título como modelo
    paragrafo_titulo_modelo = None
    for para in doc.paragraphs:
        if '{NOME_ESPECIALIDADE}' in para.text:
            paragrafo_titulo_modelo = para
            break
    
    # Remove APENAS a tabela modelo e o título com variável
    elementos_para_remover = []
    
    for elemento in doc.element.body:
        if elemento.tag.endswith('tbl'):
            # Remove tabelas
            elementos_para_remover.append(elemento)
        elif elemento.tag.endswith('p') and paragrafo_titulo_modelo:
            # Remove o parágrafo do título modelo
            if elemento == paragrafo_titulo_modelo._element:
                elementos_para_remover.append(elemento)
    
    # Remove os elementos identificados
    for elem in elementos_para_remover:
        try:
            elem.getparent().remove(elem)
        except Exception as e:
            logger.warning(f"⚠ Não foi possível remover elemento: {e}")
    
    # Agrupa por especialidade
    especialidades = sorted(list(set(d['procedimento'] for d in dados)))
    logger.info(f"→ Processando {len(especialidades)} especialidade(s): {', '.join(especialidades)}")

    total_linhas_geradas = 0
    duracao = config["duracao_atendimento_minutos"]

    for esp in especialidades:
        atendimentos_esp = [d for d in dados if d['procedimento'] == esp]
        logger.info(f"  • {esp}: {len(atendimentos_esp)} atendimento(s)")
        
        # CLONA o parágrafo do título preservando toda a formatação
        from copy import deepcopy
        if paragrafo_titulo_modelo:
            # Clona o elemento XML do título
            titulo_element = deepcopy(paragrafo_titulo_modelo._element)
            doc._element.body.append(titulo_element)
            
            # Acessa o novo parágrafo adicionado
            novo_titulo = doc.paragraphs[-1]
            
            # Substitui a variável pelo nome da especialidade
            for run in novo_titulo.runs:
                if '{NOME_ESPECIALIDADE}' in run.text:
                    run.text = run.text.replace('{NOME_ESPECIALIDADE}', esp)
        else:
            # Fallback: se não encontrou o modelo, adiciona texto simples
            titulo = doc.add_paragraph(esp)
            titulo.style = 'Normal'
        
        # CLONA A TABELA COMPLETA (preserva tudo: larguras, alturas, bordas, cores)
        nova_tabela = clonar_tabela_completa(tabela_modelo, doc)
        
        # Remove todas as linhas de dados (mantém só o cabeçalho)
        # A tabela clonada vem com todas as linhas do template
        linhas_para_remover = len(nova_tabela.rows) - 1  # Remove todas exceto cabeçalho
        for _ in range(linhas_para_remover):
            # Remove sempre a segunda linha (índice 1, pois índice 0 é o cabeçalho)
            if len(nova_tabela.rows) > 1:
                nova_tabela._element.remove(nova_tabela.rows[1]._element)

        # Preenche a tabela com dados da especialidade
        for item in atendimentos_esp:
            try:
                inicio = datetime.strptime(item['inicio'], config["formato_hora"])
                termino = (inicio + timedelta(minutes=duracao)).strftime(config["formato_hora"])
                
                # Clona a linha de dados do modelo (linha 1 da tabela modelo)
                nova_linha = clonar_linha_tabela(nova_tabela, linha_dados_modelo)
                
                # Dados da linha: [DATA, INÍCIO, TÉRMINO, EVOLUÇÃO DIÁRIA (vazio), TÉCNICO (vazio)]
                dados_linha = [
                    item['data'],
                    item['inicio'],
                    termino,
                    "",  # EVOLUÇÃO DIÁRIA
                    ""   # TÉCNICO
                ]
                
                preencher_linha_tabela(nova_linha, dados_linha)
                total_linhas_geradas += 1
                
            except ValueError as e:
                logger.error(f"✗ Erro ao processar {esp} (Tabela {item['tabela_origem']}, Linha {item['linha_origem']}): {e}")

    try:
        doc.save(caminho_destino)
        
        # Limpa o arquivo temporário
        try:
            os.unlink(temp_template_path)
        except:
            pass
        
        logger.info(f"\n{'='*60}")
        logger.info(f"✓ SUCESSO! Documento gerado:")
        logger.info(f"  • Arquivo: '{caminho_destino}'")
        logger.info(f"  • Total de linhas: {total_linhas_geradas}")
        logger.info(f"  • Especialidades: {len(especialidades)}")
        logger.info(f"  • Duração padrão: {duracao} minutos")
        logger.info(f"  • Formatação: 100% preservada do template")
        logger.info(f"{'='*60}\n")
        return True
    except Exception as e:
        logger.error(f"✗ ERRO ao salvar documento: {e}")
        # Limpa o arquivo temporário mesmo em caso de erro
        try:
            os.unlink(temp_template_path)
        except:
            pass
        return False

def gerar_config_exemplo():
    """Gera um arquivo de configuração de exemplo."""
    caminho = 'config.json'
    if not os.path.exists(caminho):
        with open(caminho, 'w', encoding='utf-8') as f:
            json.dump(CONFIG_PADRAO, f, indent=4, ensure_ascii=False)
        logger.info(f"✓ Arquivo de configuração exemplo criado: '{caminho}'")
    else:
        logger.info(f"ℹ Arquivo '{caminho}' já existe. Não será sobrescrito.")

def criar_diretorios_padrao():
    """Cria os diretórios de entrada e saída se não existirem."""
    for diretorio in ['entrada', 'saida']:
        if not os.path.exists(diretorio):
            os.makedirs(diretorio)
            logger.info(f"✓ Diretório '{diretorio}/' criado")

def main():
    """Função principal com argumentos de linha de comando."""
    logger.info("="*60)
    logger.info("GERADOR DE FOLHA DE EVOLUÇÃO TRANSDISCIPLINAR")
    logger.info("="*60 + "\n")
    
    # Garante que os diretórios existem
    criar_diretorios_padrao()
    
    # Verifica argumentos de linha de comando
    if len(sys.argv) == 3:
        arquivo_origem = sys.argv[1]
        arquivo_destino = sys.argv[2]
        logger.info(f"→ Modo: Linha de comando")
    elif len(sys.argv) == 2 and sys.argv[1] == '--gerar-config':
        gerar_config_exemplo()
        return
    else:
        # Modo padrão (hardcoded para uso pessoal)
        arquivo_origem = 'entrada/JOAO PAULO NUNES - Folha de frequência JULHO.docx'
        arquivo_destino = 'saida/Evolucao_Julho_Final.docx'
        logger.info(f"→ Modo: Execução padrão")
        logger.info(f"ℹ Dica: Use 'python main.py <origem.docx> <destino.docx>' para especificar arquivos")
        logger.info(f"ℹ Dica: Use 'python main.py --gerar-config' para criar config.json\n")
    
    # Carrega configurações
    config = carregar_configuracao()
    
    # Validação de entrada
    if not validar_arquivo_entrada(arquivo_origem):
        logger.error("\n✗ Processo interrompido devido a erros de validação.")
        sys.exit(1)
    
    # Extração de dados do cabeçalho (se configurado)
    dados_cabecalho = None
    if config.get("extrair_cabecalho_de_entrada", False):
        dados_cabecalho = extrair_dados_cabecalho(arquivo_origem)
    
    # Extração de dados das tabelas
    dados, erros, avisos = identificar_e_extrair_tabelas(arquivo_origem, config)
    
    if not dados:
        logger.error("\n✗ Processo interrompido: nenhum dado válido extraído.")
        sys.exit(1)
    
    # Geração do documento
    sucesso = gerar_word_evolucao(dados, arquivo_destino, config, dados_cabecalho)
    
    if not sucesso:
        logger.error("\n✗ Processo concluído com erros.")
        sys.exit(1)
    
    logger.info("✓ Processo concluído com sucesso!")

if __name__ == "__main__":
    main()